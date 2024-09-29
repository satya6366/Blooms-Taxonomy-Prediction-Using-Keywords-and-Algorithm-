import pandas as pd
import torch
from transformers import BertTokenizer, BertForSequenceClassification
from transformers import Trainer, TrainingArguments
from sklearn.model_selection import train_test_split
import pdfkit
from docx import Document  # Import Document class from docx module
from flask import Flask, render_template, request, make_response, send_file, redirect, url_for
import pdfkit
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


app = Flask(__name__)

# Load dataset
data = pd.read_csv('bloom_data.csv')

# Preprocessing
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
labels = data['label'].unique().tolist()
label_map = {label: idx for idx, label in enumerate(labels)}

def preprocess_function(examples):
    return tokenizer(examples['text'], truncation=True, padding=True)

data['label'] = data['label'].map(label_map)

class BloomDataset(torch.utils.data.Dataset):
    def __init__(self, encodings, labels):
        self.encodings = encodings
        self.labels = labels

    def __getitem__(self, idx):
        item = {key: torch.tensor(val[idx]) for key, val in self.encodings.items()}
        item['labels'] = torch.tensor(self.labels[idx])
        return item

    def __len__(self):
        return len(self.labels)

# Split dataset
train_texts, val_texts, train_labels, val_labels = train_test_split(data['text'], data['label'], test_size=0.2)

# Tokenize
train_encodings = tokenizer(train_texts.tolist(), truncation=True, padding=True)
val_encodings = tokenizer(val_texts.tolist(), truncation=True, padding=True)

train_dataset = BloomDataset(train_encodings, train_labels.tolist())
val_dataset = BloomDataset(val_encodings, val_labels.tolist())

# Model
model = BertForSequenceClassification.from_pretrained('bert-base-uncased', num_labels=len(labels))

# Training
training_args = TrainingArguments(
    output_dir='./results',
    num_train_epochs=3,
    per_device_train_batch_size=16,
    per_device_eval_batch_size=16,
    warmup_steps=500,
    weight_decay=0.01,
    logging_dir='./logs',
)

trainer = Trainer(
    model=model,
    args=training_args,
    train_dataset=train_dataset,
    eval_dataset=val_dataset,
)

trainer.train()

# Function to classify new text
def classify_text(text):
    inputs = tokenizer(text, return_tensors='pt', truncation=True, padding=True)
    outputs = model(**inputs)
    predictions = torch.argmax(outputs.logits, dim=1)
    return labels[predictions.item()]

@app.route('/')
def index():
    return render_template('index.html', level1=None)

@app.route('/qp')
def qp():
    return render_template('qp.html')

@app.route('/sem')
def sem():
    return render_template('sem.html', bloom_level=None)

@app.route('/analyze', methods=['POST'])
def analyze():
    question1 = request.form.get('question1')
    print("Received question:", question1)  # Debugging line
    if question1:
        level1 = analyze_blooms_taxonomy(question1)
        print("Analyzed level:", level1)  # Debugging line
    else:
        level1 = "No question provided"
    return render_template('index.html', level1=level1, question1=question1)



@app.route('/qpanalyze', methods=['POST'])
def qpanalyze():
    data = {
        'course': request.form.get('course'),
        'course_code': request.form.get('course_code'),
        'semester': request.form.get('semester'),
        'max_marks': request.form.get('max_marks'),
        'date': request.form.get('date'),
        'duration': request.form.get('duration'),
        'co': {f'co_{i}{j}': request.form.get(f'co_{i}{j}') for i in range(1, 9) for j in ['a', 'b']},
        'po': {f'po_{i}{j}': request.form.get(f'po_{i}{j}') for i in range(1, 9) for j in ['a', 'b']},
        'questions': {f'question{i}{j}': request.form.get(f'question{i}{j}') for i in range(1, 9) for j in ['a', 'b']},
        'marks': {f'marks{i}': request.form.get(f'marks{i}') for i in range(1, 9)}
    }
    
    # Analyze Bloom's levels
    blooms_levels = {}
    for i in range(1, 9):
        for part in ['a', 'b']:
            question_key = f'question{i}{part}'
            bloom_key = f'level{i}{part}'
            if question_key in data['questions']:
                question = data['questions'][question_key]
                level = analyze_blooms(question)
                blooms_levels[bloom_key] = level
    
    # Flatten the data dictionary for template access
    flattened_data = {**data, **data['co'], **data['po'], **data['questions'], **data['marks'], **blooms_levels}
    
    return render_template('analysis.html', data=flattened_data)

@app.route('/download', methods=['POST'])
@app.route('/download', methods=['POST'])
def download():
    data = request.form.to_dict()
    file_type = request.form['file_type']

    if file_type == 'pdf':
        path_to_wkhtmltopdf = '/usr/local/bin/wkhtmltopdf'  # Update this path based on your system
        config = pdfkit.configuration(wkhtmltopdf=path_to_wkhtmltopdf)
        rendered = render_template('analysis.html', data=data)
        pdf = pdfkit.from_string(rendered, False, configuration=config)

        response = make_response(pdf)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = 'attachment; filename=question_paper_analysis.pdf'
        return response

    elif file_type == 'word':
        document = Document()
        document.styles['Normal'].font.name = 'Times New Roman'
        document.styles['Normal'].font.size = Pt(12)

        # Add header
        header = document.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = (
            "DAYANANDA SAGAR COLLEGE OF ENGINEERING\n"
            "[An Autonomous Institution Affiliated to VTU, Belagavi, Accredited by NBA (CSE, ECE, EEE, ISE, ME) and NAAC with 'A' Grade]\n"
            "Shavige Malleshwara Hills, Kumaraswamy Layout, Bengaluru-560111"
        )
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

         # Add logo image
        logo_path = os.path.join('static', 'img', 'logodsce.jpg')
        if os.path.exists(logo_path):
            header_paragraph = header.add_paragraph()
            run = header_paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1))
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        document.add_heading('Question Paper Analysis', level=1)

        # Add Course Information
        document.add_heading('Course Information', level=2)
        document.add_paragraph(f"Course: {data['course']}")
        document.add_paragraph(f"Course Code: {data['course_code']}")
        document.add_paragraph(f"Semester: {data['semester']}")
        document.add_paragraph(f"Maximum Marks: {data['max_marks']}")
        document.add_paragraph(f"Date: {data['date']}")
        document.add_paragraph(f"Duration: {data['duration']}")

        # Add Details Table
        document.add_heading('Details', level=2)
        table = document.add_table(rows=4, cols=17)
        headers = ['Question No.', '1 a', '1 b', '2 a', '2 b', '3 a', '3 b', '4 a', '4 b', '5 a', '5 b', '6 a', '6 b', '7 a', '7 b', '8 a', '8 b']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        row_cells = table.rows[1].cells
        row_cells[0].text = 'Course Outcome'
        for i in range(1, 9):
            row_cells[i*2-1].text = data.get(f'co_{i}a', '')
            row_cells[i*2].text = data.get(f'co_{i}b', '')

        row_cells = table.rows[2].cells
        row_cells[0].text = 'Program Outcome'
        for i in range(1, 9):
            row_cells[i*2-1].text = data.get(f'po_{i}a', '')
            row_cells[i*2].text = data.get(f'po_{i}b', '')

        row_cells = table.rows[3].cells
        row_cells[0].text = 'Blooms Level'
        for i in range(1, 9):
            row_cells[i*2-1].text = data.get(f'level{i}a', '')
            row_cells[i*2].text = data.get(f'level{i}b', '')

        # Add Questions
        document.add_heading('Questions', level=2)
        for i in range(1, 9):
            document.add_heading(f'Question {i} for 10 Marks', level=3)
            document.add_paragraph(f"{i}a) {data.get(f'question{i}a', '')}")
            document.add_paragraph(f"{i}b) {data.get(f'question{i}b', '')}")
            document.add_paragraph(f"Marks: {data.get(f'marks{i}', '')}")
            if i < 8:
                document.add_paragraph("(OR)")

        file_path = '/tmp/question_paper_analysis.docx'
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        document.save(file_path)
        return send_file(file_path, as_attachment=True, download_name='question_paper_analysis.docx')

    return redirect(url_for('qp'))



    
@app.route('/analyze_sem', methods=['POST'])
def analyze_sem():
    semester = request.form.get('semester')
    question = request.form.get('question')
    if not semester or not question:
        return render_template('sem.html', bloom_level="Please provide both semester and question.")

    semester = int(semester)
    filtered_data = data[data['semester'] == semester]
    
    if filtered_data.empty:
        return render_template('sem.html', bloom_level=f"No data available for semester {semester}.")

    bloom_level = classify_text(question)
    return render_template('sem.html', bloom_level=f"The Bloom's level for the given {question} is: level {bloom_level}")

def analyze_blooms_taxonomy(question1):
    keywords_to_levels = {
        'Remembering level 1': ['define', 'describe', 'find', 'how', 'list', 'name', 'what', 'where', 'which', 'why', 'draw', 'write'],
        'Understanding level 2': ['compare', 'demonstrate', 'discuss', 'distinguish', 'explain', 'illustrate', 'outline', 'show', 'summarize'],
        'Applying level 3': ['compute', 'develop','analyse','analyze', 'identify', 'make use of', 'select', 'solve', 'utilize', 'use', 'draw', 'illustrate', 'classify', 'solve', 'categorize'],
        'Analyzing level 4': ['classify', 'characterize', 'categorize', 'compare', 'derive', 'distinguish', 'examine', 'inference', 'organize', 'simplify', 'test for', 'identify', 'investigate'],
        'Evaluating level 5': ['assess', 'choose', 'compare', 'decide', 'determine', 'estimate', 'evaluate', 'explain', 'interpret', 'justify', 'measure', 'prioritize', 'prove', 'rate', 'recommend'],
        'Creating level 6': ['build', 'compose', 'construct', 'create', 'design', 'develop', 'discuss', 'elaborate', 'estimate', 'formulate', 'improve', 'maximize', 'modify', 'predict', 'invent']
    }

    # Normalize the input question
    question_lower = question1.lower()

    # Iterate through each level and its keywords to find a match
    for level1, keywords in keywords_to_levels.items():
        for keyword in keywords:
            if keyword in question_lower:
                return level1

    # If no keywords match, return "Uncategorized"
    return "Uncategorized"

def analyze_blooms(question):
    keywords_to_levels = {
        '1': ['define', 'describe', 'find', 'how', 'list', 'name', 'what', 'where', 'which', 'why', 'draw', 'write'],
        '2': ['compare', 'demonstrate', 'discuss', 'distinguish', 'explain', 'illustrate', 'outline', 'show', 'summarize'],
        '3': ['compute', 'analyse','analyze','develop', 'identify', 'make use of', 'select', 'solve', 'utilize', 'use', 'draw', 'illustrate', 'classify', 'solve', 'categorize'],
        '4': ['classify', 'characterize', 'categorize', 'compare', 'derive', 'distinguish', 'examine', 'inference', 'organize', 'simplify', 'test for', 'identify', 'investigate'],
        '5': ['assess', 'choose', 'compare', 'decide', 'determine', 'estimate', 'evaluate', 'explain', 'interpret', 'justify', 'measure', 'prioritize', 'prove', 'rate', 'recommend'],
        '6': ['build', 'compose', 'construct', 'create', 'design', 'develop', 'discuss', 'elaborate', 'estimate', 'formulate', 'improve', 'maximize', 'modify', 'predict', 'invent']
    }

    question_lower = question.lower()
    for level, keywords in keywords_to_levels.items():
        for keyword in keywords:
            if keyword in question_lower:
                return level

    return 'UN'

if __name__ == '__main__':
    app.run(debug=True)
